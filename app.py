import streamlit as st
from docx import Document
from docx.shared import Inches
from PIL import Image
import io
import tempfile
from pdf2image import convert_from_bytes
import pdfplumber

st.set_page_config(page_title="Generador de Informes", layout="centered")
st.title("Generador de Informes Periciales")

# Carga de texto
texto_encargo = st.text_area("Pega el texto del encargo", height=300)

# Subida de imágenes para informe
fotos = st.file_uploader("Selecciona imágenes del informe (puedes subir varias)", type=["jpg", "jpeg", "png"], accept_multiple_files=True)

# Subida del archivo del catastro
catastro_img = st.file_uploader("Imagen del Catastro (JPG o PNG)", type=["jpg", "jpeg", "png"])
catastro_pdf = st.file_uploader("O Catastro en PDF", type=["pdf"])

# Botón para generar informe
if st.button("Generar informe") and texto_encargo:
    rep = {"{{IMG_CATASTRO}}": None, "{{PROVINCIA_CATASTRO}}": "XXX"}

    # --- PARSEO DEL TEXTO PEGADO ---
    for line in texto_encargo.splitlines():
        if ":" in line:
            key, value = line.split(":", 1)
            rep[f"{{{key.strip().upper()}}}"] = value.strip()

    # --- IMAGEN CATÁSTRO ---
    if catastro_img:
        rep["{{IMG_CATASTRO}}"] = Image.open(catastro_img)
    elif catastro_pdf:
        try:
            img = convert_from_bytes(catastro_pdf.read(), first_page=1, last_page=1, dpi=200)[0]
            rep["{{IMG_CATASTRO}}"] = img
            catastro_pdf.seek(0)
            with pdfplumber.open(catastro_pdf) as pdf:
                texto_pdf = pdf.pages[0].extract_text()
                for line in texto_pdf.split('\n'):
                    if "Provincia:" in line:
                        rep["{{PROVINCIA_CATASTRO}}"] = line.split(":")[1].strip().upper()
        except Exception as e:
            st.warning("No se pudo extraer la imagen del catastro desde el PDF. Suba una imagen si es necesario.")

    # --- CARGA PLANTILLA ---
    plantilla = "PLANTILLA BASE - V3 generador.docx"
    if "15 - ASIST.JURIDICA" in texto_encargo:
        plantilla = "PLANTILLA BASE JURIDICO - V3 generador.docx"
    doc = Document(plantilla)

    # --- REEMPLAZO TEXTO ---
    for p in doc.paragraphs:
        for key, val in rep.items():
            if isinstance(val, str) and key in p.text:
                inline = p.runs
                for i in range(len(inline)):
                    if key in inline[i].text:
                        inline[i].text = inline[i].text.replace(key, val)

    # --- IMAGEN CATASTRO ---
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                if "{{IMG_CATASTRO}}" in cell.text:
                    cell.text = ""
                    if rep["{{IMG_CATASTRO}}"]:
                        cell.paragraphs[0].add_run().add_picture(rep["{{IMG_CATASTRO}}"], width=Inches(2.2))

    # --- AÑADIR IMÁGENES AL FINAL ---
    if fotos:
        doc.add_page_break()
        doc.add_paragraph("Reportaje fotográfico")
        for foto in fotos:
            doc.add_picture(foto, width=Inches(4.5))
            doc.add_paragraph(foto.name)

    # --- GUARDAR Y DESCARGAR ---
    buffer = io.BytesIO()
    nombre_archivo = f"{rep.get('{{EXPEDIENTE}}', 'SIN_EXP')}.docx"
    doc.save(buffer)
    buffer.seek(0)
    st.download_button("Descargar informe Word", buffer, file_name=nombre_archivo, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")